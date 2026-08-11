import sys, os, io
sys.path.insert(0, '.')
os.environ['FLASK_ENV'] = 'development'

from unittest.mock import patch
import app as app_module
from ai.base import BaseAIProvider, AIResult
from docx import Document


class FakeProvider(BaseAIProvider):
    name = "fake"
    model = "fake-model-1"
    def __init__(self, canned_text, configured=True):
        self.canned_text = canned_text
        self.configured = configured
    def is_configured(self):
        return self.configured
    def _call(self, system_prompt, user_prompt, max_tokens):
        return AIResult(text=self.canned_text, provider=self.name, model=self.model, latency_ms=0)


def _client():
    app_module.app.config['TESTING'] = True
    return app_module.app.test_client()


def _setup_session_with_data(client):
    sess = client.post('/api/sessions', json={
        'name': 'Faz4 Route Test', 'tester_name': 'Test Tester', 'target_url': 'https://example.com'
    }).get_json()
    sid = sess['id']
    client.post(f'/api/sessions/{sid}/results', json={'test_id': 'WSTG-ATHN-01', 'status': 'done'})
    client.post(f'/api/sessions/{sid}/notes', json={
        'title': 'Login SQLi', 'content': "' OR 1=1-- ile bypass", 'severity': 'critical',
        'test_id': 'WSTG-ATHN-01', 'cwe_id': 'CWE-89', 'cwe_name': 'SQL Injection',
        'cvss_vector': 'AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H'
    })
    return sid


def test_report_data_endpoint():
    client = _client()
    sid = _setup_session_with_data(client)
    resp = client.get(f'/api/sessions/{sid}/report/data')
    assert resp.status_code == 200
    data = resp.get_json()
    assert data['stats']['total_findings'] == 1
    assert data['stats']['severity_counts']['critical'] == 1
    assert data['findings'][0]['cvss_score'] == 9.8
    print("OK: /report/data zenginleştirilmiş rapor verisini doğru döndürüyor")


def test_report_summary_endpoint():
    client = _client()
    sid = _setup_session_with_data(client)
    canned = "Bu ortamda kritik bir SQL Injection zafiyeti bulunmuştur."
    with patch.object(app_module, 'get_ai_provider', return_value=FakeProvider(canned)):
        resp = client.post(f'/api/sessions/{sid}/report/summary', json={})
    assert resp.status_code == 200
    body = resp.get_json()
    assert body['summary'] == canned
    assert body['provider'] == 'fake'

    logs = client.get(f'/api/ai/logs?session_id={sid}&purpose=report_generation').get_json()
    assert len(logs) == 1 and logs[0]['success'] is True
    print("OK: /report/summary AI özet taslağı üretiyor ve loglanıyor")


def test_report_download_markdown():
    client = _client()
    sid = _setup_session_with_data(client)
    resp = client.post(f'/api/sessions/{sid}/report/download', json={
        'format': 'md', 'executive_summary': 'Düzenlenmiş özet metni.'
    })
    assert resp.status_code == 200
    assert resp.mimetype == 'text/markdown'
    text = resp.get_data(as_text=True)
    assert 'Düzenlenmiş özet metni.' in text
    assert 'Login SQLi' in text
    print("OK: /report/download markdown formatında düzenlenmiş özetle birlikte gerçek bir dosya döndürüyor")


def test_report_download_docx():
    client = _client()
    sid = _setup_session_with_data(client)
    resp = client.post(f'/api/sessions/{sid}/report/download', json={
        'format': 'docx', 'executive_summary': 'Word raporu için özet.'
    })
    assert resp.status_code == 200
    assert 'wordprocessingml' in resp.mimetype
    doc = Document(io.BytesIO(resp.data))
    all_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        c.text for t in doc.tables for r in t.rows for c in r.cells
    )
    assert 'Word raporu için özet.' in all_text
    assert 'Login SQLi' in all_text
    assert 'CWE-89' in all_text
    print("OK: /report/download geçerli, içeriği doğru bir .docx dosyası döndürüyor")


def test_report_download_without_summary_still_works():
    client = _client()
    sid = _setup_session_with_data(client)
    resp = client.post(f'/api/sessions/{sid}/report/download', json={'format': 'md'})
    assert resp.status_code == 200
    assert 'Yönetici Özeti' not in resp.get_data(as_text=True)
    print("OK: yönetici özeti verilmeden de rapor indirilebiliyor (o bölüm atlanıyor)")


def test_report_download_invalid_format_rejected():
    client = _client()
    sid = _setup_session_with_data(client)
    resp = client.post(f'/api/sessions/{sid}/report/download', json={'format': 'pdf'})
    assert resp.status_code == 400
    print("OK: geçersiz format 400 ile reddediliyor")


if __name__ == '__main__':
    test_report_data_endpoint()
    test_report_summary_endpoint()
    test_report_download_markdown()
    test_report_download_docx()
    test_report_download_without_summary_still_works()
    test_report_download_invalid_format_rejected()
    print("\nTüm Faz 4 rapor endpoint testleri geçti.")

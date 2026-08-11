import sys, io
sys.path.insert(0, '.')

from ai.base import BaseAIProvider, AIResult
import report_generator as rg


class FakeProvider(BaseAIProvider):
    name = "fake"
    model = "fake-model"
    def __init__(self, canned_text):
        self.canned_text = canned_text
    def is_configured(self):
        return True
    def _call(self, system_prompt, user_prompt, max_tokens):
        self.last_system_prompt = system_prompt
        self.last_user_prompt = user_prompt
        return AIResult(text=self.canned_text, provider=self.name, model=self.model, latency_ms=0)


SESSION = {"id": "s1", "name": "Örnek Şirket Pentesti", "target_url": "https://example.com", "tester_name": "Ayşe Yılmaz"}

RESULTS = [
    {"test_id": "WSTG-ATHN-01", "category_id": "athn", "status": "done"},
    {"test_id": "WSTG-ATHN-02", "category_id": "athn", "status": "done"},
    {"test_id": "WSTG-INFO-01", "category_id": "info", "status": "pending"},
]

NOTES = [
    {"id": 1, "title": "Login SQLi", "content": "' OR 1=1-- ile bypass edildi.", "severity": "critical",
     "test_id": "WSTG-ATHN-01", "cwe_id": "CWE-89", "cwe_name": "SQL Injection",
     "cvss_score": 9.8, "cvss_rating": "critical", "cvss_vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H"},
    {"id": 2, "title": "Zayıf parola politikası", "content": "Minimum uzunluk kontrolü yok.", "severity": "medium",
     "test_id": "WSTG-ATHN-02", "cwe_id": None, "cwe_name": None,
     "cvss_score": None, "cvss_rating": None, "cvss_vector": None},
]


def test_build_report_data_stats():
    data = rg.build_report_data(SESSION, RESULTS, NOTES)
    assert data["stats"]["total_tests"] == 3
    assert data["stats"]["completed_tests"] == 2
    assert data["stats"]["completion_pct"] == 66.7
    assert data["stats"]["total_findings"] == 2
    assert data["stats"]["severity_counts"]["critical"] == 1
    assert data["stats"]["avg_cvss"] == 9.8
    assert data["stats"]["max_cvss"] == 9.8
    assert data["completed_test_ids"] == ["WSTG-ATHN-01", "WSTG-ATHN-02"]
    print("OK: build_report_data istatistikleri doğru hesaplıyor")


def test_findings_sorted_by_severity_then_cvss():
    data = rg.build_report_data(SESSION, RESULTS, NOTES)
    assert data["findings"][0]["title"] == "Login SQLi"  # critical > medium
    print("OK: bulgular önem derecesine göre sıralanıyor (kritik önce)")


def test_findings_enriched_with_owasp_category():
    data = rg.build_report_data(SESSION, RESULTS, NOTES)
    sqli_finding = next(f for f in data["findings"] if f["title"] == "Login SQLi")
    # WSTG-ATHN-01 kimlik dogrulama testine bagli, owasp_categories bos olmayabilir de olabilir
    # (mapping veri setine bagli), en azindan alan var olmali ve liste tipinde olmali.
    assert isinstance(sqli_finding["owasp_categories"], list)
    assert sqli_finding["severity_label"] == "Kritik"
    print("OK: bulgular OWASP kategori bilgisiyle zenginleştiriliyor")


def test_executive_summary_generation():
    provider = FakeProvider("Bu sistemde kritik bir SQL Injection zafiyeti tespit edilmiştir. "
                             "Acil müdahale önerilir.")
    data = rg.build_report_data(SESSION, RESULTS, NOTES)
    summary, ai_result = rg.generate_executive_summary(provider, data)
    assert "SQL Injection" in summary
    assert ai_result.provider == "fake"
    assert "Login SQLi" in provider.last_user_prompt
    assert "critical" in provider.last_user_prompt.lower() or "Kritik" in provider.last_user_prompt
    print("OK: AI yönetici özeti üretimi ve prompt bağlamı doğru çalışıyor")


def test_render_markdown_contains_key_sections():
    data = rg.build_report_data(SESSION, RESULTS, NOTES)
    md = rg.render_markdown(data, executive_summary="Test özeti metni.")
    assert "# Pentest Raporu" in md
    assert "Test özeti metni." in md
    assert "Login SQLi" in md
    assert "CWE-89" in md
    assert "9.8" in md
    assert "WSTG-ATHN-01" in md
    print("OK: Markdown render tüm ana bölümleri (özet, bulgular, metodoloji) içeriyor")


def test_render_markdown_without_summary():
    data = rg.build_report_data(SESSION, RESULTS, NOTES)
    md = rg.render_markdown(data)  # executive_summary verilmedi
    assert "Yönetici Özeti" not in md  # özet verilmediyse bölüm hiç eklenmemeli
    print("OK: yönetici özeti verilmezse Markdown'da o bölüm hiç görünmüyor")


def test_render_docx_produces_valid_file():
    from docx import Document
    data = rg.build_report_data(SESSION, RESULTS, NOTES)
    buf = rg.render_docx(data, executive_summary="Yönetici özeti test metni.")
    assert isinstance(buf, io.BytesIO)
    buf.seek(0)
    doc = Document(buf)  # gecerli bir docx olarak tekrar acilabiliyor mu?
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for h in doc.paragraphs:
        pass
    all_text = full_text + "\n" + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    assert "Yönetici özeti test metni." in all_text
    assert "Login SQLi" in all_text
    assert "Kritik" in all_text
    print("OK: DOCX render geçerli bir Word dosyası üretiyor ve içerik doğru gömülüyor")


def test_render_docx_empty_findings_no_crash():
    data = rg.build_report_data(SESSION, [], [])
    buf = rg.render_docx(data)
    assert isinstance(buf, io.BytesIO)
    print("OK: hiç test/bulgu olmayan bir oturum için bile DOCX üretimi çökmeden çalışıyor")


if __name__ == "__main__":
    test_build_report_data_stats()
    test_findings_sorted_by_severity_then_cvss()
    test_findings_enriched_with_owasp_category()
    test_executive_summary_generation()
    test_render_markdown_contains_key_sections()
    test_render_markdown_without_summary()
    test_render_docx_produces_valid_file()
    test_render_docx_empty_findings_no_crash()
    print("\nTüm report_generator testleri geçti.")

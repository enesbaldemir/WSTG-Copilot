"""
Otomatik Rapor Uretimi (Faz 4).

Bir oturumun (session) toplanan bulgularindan (Not Defteri kayitlari,
CVSS/CWE ile zenginlestirilmis) ve tamamlanan WSTG testlerinden hem
JSON hem de indirilebilir Markdown/DOCX formatinda bir pentest raporu
uretir.

Onemli tasarim karari: AI sadece YONETICI OZETI (executive summary)
metnini yazar ve bu metin dogrudan rapora gomulmez -- once frontend'de
duzenlenebilir bir onizlemede gosterilir, pentester onaylayip/duzenleyip
"Indir" dedikten SONRA rapora islenir. Boylece musteriye giden nihai
belgede AI'nin denetimsiz/gozden gecirilmemis bir metni bulunmaz.
"""

import io
from collections import Counter
from datetime import datetime

import mapping as mapping_lib

SEVERITY_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4, None: 5}

SEVERITY_COLORS_RGB = {
    "critical": "C0392B",
    "high": "E67E22",
    "medium": "F1C40F",
    "low": "27AE60",
    "info": "7F8C8D",
}

SEVERITY_LABEL_TR = {"critical": "Kritik", "high": "Yüksek", "medium": "Orta", "low": "Düşük", "info": "Bilgi"}
SEVERITY_LABEL_EN = {"critical": "Critical", "high": "High", "medium": "Medium", "low": "Low", "info": "Info"}


def _severity_label(sev, lang):
    table = SEVERITY_LABEL_TR if lang == "tr" else SEVERITY_LABEL_EN
    return table.get(sev, sev or "-")


def build_report_data(session, results, notes, lang="tr"):
    """
    session: Session.to_dict()
    results: [TestResult.to_dict(), ...]
    notes:   [Note.to_dict(), ...]
    """
    total_tests = len(results)
    completed_results = [r for r in results if r.get("status") != "pending"]
    completed_tests = len(completed_results)
    completion_pct = round((completed_tests / total_tests) * 100, 1) if total_tests else 0.0

    findings = sorted(
        notes,
        key=lambda n: (SEVERITY_ORDER.get(n.get("severity"), 5), -(n.get("cvss_score") or 0))
    )

    severity_counts = Counter(n.get("severity") or "info" for n in findings)
    cvss_scores = [n["cvss_score"] for n in findings if n.get("cvss_score") is not None]
    avg_cvss = round(sum(cvss_scores) / len(cvss_scores), 1) if cvss_scores else None
    max_cvss = max(cvss_scores) if cvss_scores else None

    # Her bulguya, varsa ilişkili OWASP Top 10 kategorisini de ekleyelim (rapor bağlamı için).
    enriched_findings = []
    for n in findings:
        owasp_matches = mapping_lib.get_mapping_for_test(n["test_id"], lang) if n.get("test_id") else []
        enriched_findings.append({
            **n,
            "severity_label": _severity_label(n.get("severity"), lang),
            "owasp_categories": [f"{m['owasp_id']} {m['owasp_title']}" for m in owasp_matches],
        })

    completed_by_category = {}
    for r in completed_results:
        cat = r.get("category_id") or "?"
        completed_by_category.setdefault(cat, []).append(r["test_id"])

    return {
        "session": session,
        "generated_at": datetime.utcnow().isoformat(),
        "stats": {
            "total_tests": total_tests,
            "completed_tests": completed_tests,
            "completion_pct": completion_pct,
            "total_findings": len(findings),
            "severity_counts": dict(severity_counts),
            "avg_cvss": avg_cvss,
            "max_cvss": max_cvss,
        },
        "findings": enriched_findings,
        "completed_test_ids": sorted(r["test_id"] for r in completed_results),
    }


# ---------------------------------------------------------------------------
# AI Yönetici Özeti
# ---------------------------------------------------------------------------

def build_executive_summary_prompt(report_data, lang="tr"):
    stats = report_data["stats"]
    session = report_data["session"]
    top_findings = report_data["findings"][:8]

    findings_lines = []
    for f in top_findings:
        cwe = f" ({f['cwe_id']})" if f.get("cwe_id") else ""
        findings_lines.append(f"- [{f['severity_label']}] {f.get('title') or '(başlıksız)'}{cwe}")
    findings_block = "\n".join(findings_lines) if findings_lines else "(Kayıtlı bulgu yok.)"

    system_prompt = (
        "Sen deneyimli bir pentest raporu yazarısın. Sana bir güvenlik testi oturumunun "
        "istatistikleri ve bulgu listesi verilecek. Görevin, teknik olmayan bir yöneticinin de "
        "anlayabileceği, 150-250 kelimelik bir YÖNETİCİ ÖZETİ (executive summary) yazmak: genel risk "
        "durumu, en kritik bulgular, iş etkisi ve genel tavsiye. Teknik jargondan kaçın, CVSS/CWE "
        "kodlarını tekrarlama. SADECE özet metnini yaz — başlık, madde işareti, markdown biçimlendirmesi "
        "veya ek açıklama EKLEME, düz paragraf(lar) halinde yaz."
    )
    user_prompt = (
        f"Hedef: {session.get('target_url') or session.get('name')}\n"
        f"Tamamlanan test: {stats['completed_tests']}/{stats['total_tests']} (%{stats['completion_pct']})\n"
        f"Toplam bulgu: {stats['total_findings']}\n"
        f"Önem derecesi dağılımı: {stats['severity_counts']}\n"
        f"Ortalama CVSS: {stats['avg_cvss']}, En yüksek CVSS: {stats['max_cvss']}\n\n"
        f"Öne çıkan bulgular:\n{findings_block}\n"
    )
    return system_prompt, user_prompt


def generate_executive_summary(provider, report_data, lang="tr", max_tokens=1000):
    system_prompt, user_prompt = build_executive_summary_prompt(report_data, lang)
    ai_result = provider.chat(system_prompt=system_prompt, user_prompt=user_prompt, max_tokens=max_tokens)
    return ai_result.text.strip(), ai_result


# ---------------------------------------------------------------------------
# Markdown render
# ---------------------------------------------------------------------------

def render_markdown(report_data, executive_summary=None, lang="tr"):
    s = report_data["session"]
    stats = report_data["stats"]
    lines = []
    lines.append(f"# Pentest Raporu — {s.get('name') or s.get('target_url') or ''}")
    lines.append("")
    lines.append(f"- **Hedef:** {s.get('target_url') or '-'}")
    lines.append(f"- **Test Uzmanı:** {s.get('tester_name') or '-'}")
    lines.append(f"- **Oluşturulma Tarihi:** {report_data['generated_at'][:10]}")
    lines.append(f"- **Tamamlanan Testler:** {stats['completed_tests']}/{stats['total_tests']} (%{stats['completion_pct']})")
    lines.append("")

    if executive_summary:
        lines.append("## Yönetici Özeti" if lang == "tr" else "## Executive Summary")
        lines.append("")
        lines.append(executive_summary)
        lines.append("")

    lines.append("## Bulgu Özeti" if lang == "tr" else "## Findings Summary")
    lines.append("")
    lines.append("| Önem Derecesi | Sayı |" if lang == "tr" else "| Severity | Count |")
    lines.append("|---|---|")
    for sev in ["critical", "high", "medium", "low", "info"]:
        count = stats["severity_counts"].get(sev, 0)
        if count:
            lines.append(f"| {_severity_label(sev, lang)} | {count} |")
    lines.append("")

    lines.append("## Bulgular" if lang == "tr" else "## Findings")
    lines.append("")
    if not report_data["findings"]:
        lines.append("_Kayıtlı bulgu yok._" if lang == "tr" else "_No findings recorded._")
    for f in report_data["findings"]:
        lines.append(f"### [{f['severity_label']}] {f.get('title') or '(başlıksız)'}")
        meta = []
        if f.get("test_id"):
            meta.append(f"WSTG: `{f['test_id']}`")
        if f.get("cwe_id"):
            meta.append(f"CWE: `{f['cwe_id']}`" + (f" ({f['cwe_name']})" if f.get("cwe_name") else ""))
        if f.get("cvss_score") is not None:
            meta.append(f"CVSS: **{f['cvss_score']}** ({f.get('cvss_rating')}) — `{f.get('cvss_vector')}`")
        if f.get("owasp_categories"):
            meta.append("OWASP: " + ", ".join(f["owasp_categories"]))
        if meta:
            lines.append(" · ".join(meta))
        lines.append("")
        lines.append(f.get("content") or "")
        lines.append("")

    lines.append("## Metodoloji — Tamamlanan Testler" if lang == "tr" else "## Methodology — Completed Tests")
    lines.append("")
    for tid in report_data["completed_test_ids"]:
        lines.append(f"- {tid}")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# DOCX render
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def render_docx(report_data, executive_summary=None, lang="tr", output_path=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    s = report_data["session"]
    stats = report_data["stats"]
    doc = Document()

    # ---- Başlık sayfası ----
    title = doc.add_heading("Penetrasyon Testi Raporu" if lang == "tr" else "Penetration Test Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(s.get("name") or s.get("target_url") or "")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)

    meta_table = doc.add_table(rows=0, cols=2)
    meta_rows = [
        ("Hedef" if lang == "tr" else "Target", s.get("target_url") or "-"),
        ("Test Uzmanı" if lang == "tr" else "Tester", s.get("tester_name") or "-"),
        ("Tarih" if lang == "tr" else "Date", report_data["generated_at"][:10]),
        ("Tamamlanan Testler" if lang == "tr" else "Completed Tests",
         f"{stats['completed_tests']}/{stats['total_tests']} (%{stats['completion_pct']})"),
        ("Toplam Bulgu" if lang == "tr" else "Total Findings", str(stats["total_findings"])),
    ]
    for k, v in meta_rows:
        row = meta_table.add_row().cells
        row[0].text = k
        row[0].paragraphs[0].runs[0].font.bold = True
        row[1].text = str(v)

    doc.add_page_break()

    # ---- Yönetici Özeti ----
    if executive_summary:
        doc.add_heading("Yönetici Özeti" if lang == "tr" else "Executive Summary", level=1)
        doc.add_paragraph(executive_summary)

    # ---- Bulgu Özeti tablosu ----
    doc.add_heading("Bulgu Özeti" if lang == "tr" else "Findings Summary", level=1)
    sum_table = doc.add_table(rows=1, cols=2)
    sum_table.style = 'Light Grid Accent 1'
    hdr = sum_table.rows[0].cells
    hdr[0].text = "Önem Derecesi" if lang == "tr" else "Severity"
    hdr[1].text = "Sayı" if lang == "tr" else "Count"
    for sev in ["critical", "high", "medium", "low", "info"]:
        count = stats["severity_counts"].get(sev, 0)
        if not count:
            continue
        row = sum_table.add_row().cells
        row[0].text = _severity_label(sev, lang)
        row[1].text = str(count)
        _set_cell_shading(row[0], SEVERITY_COLORS_RGB.get(sev, "FFFFFF"))
        row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_page_break()

    # ---- Bulgular ----
    doc.add_heading("Bulgular" if lang == "tr" else "Findings", level=1)
    if not report_data["findings"]:
        doc.add_paragraph("Kayıtlı bulgu yok." if lang == "tr" else "No findings recorded.")

    for f in report_data["findings"]:
        h = doc.add_heading(level=2)
        run = h.add_run(f"[{f['severity_label']}] {f.get('title') or ('(başlıksız)' if lang=='tr' else '(untitled)')}")
        run.font.color.rgb = RGBColor.from_string(SEVERITY_COLORS_RGB.get(f.get("severity"), "000000"))

        meta_parts = []
        if f.get("test_id"):
            meta_parts.append(f"WSTG: {f['test_id']}")
        if f.get("cwe_id"):
            meta_parts.append(f"CWE: {f['cwe_id']}" + (f" ({f['cwe_name']})" if f.get("cwe_name") else ""))
        if f.get("cvss_score") is not None:
            meta_parts.append(f"CVSS: {f['cvss_score']} ({f.get('cvss_rating')})")
        if f.get("owasp_categories"):
            meta_parts.append("OWASP: " + ", ".join(f["owasp_categories"]))
        if meta_parts:
            meta_p = doc.add_paragraph(" · ".join(meta_parts))
            meta_p.runs[0].font.italic = True
            meta_p.runs[0].font.size = Pt(9)

        doc.add_paragraph(f.get("content") or "")

    # ---- Metodoloji ----
    doc.add_page_break()
    doc.add_heading("Metodoloji — Tamamlanan Testler" if lang == "tr" else "Methodology — Completed Tests", level=1)
    for tid in report_data["completed_test_ids"]:
        doc.add_paragraph(tid, style="List Bullet")

    if output_path:
        doc.save(output_path)
        return output_path

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
